import csv
from io import BytesIO
from decimal import Decimal
from datetime import datetime

from django.core.files.base import ContentFile
from django.http import HttpResponse
from django.db.models import Sum
from django.db.models.functions import TruncMonth
from rest_framework import permissions, response, viewsets
from rest_framework.decorators import action
from rest_framework.exceptions import PermissionDenied
from rest_framework.views import APIView

from ged.models import Document, DocumentFolder
from .models import FinanceAccount, FinanceCategory, FinanceDocument, FinanceInvoice, FinanceInvoiceLine, FinanceTransaction
from .serializers import (
    FinanceAccountSerializer,
    FinanceCategorySerializer,
    FinanceDocumentSerializer,
    FinanceInvoiceLineSerializer,
    FinanceInvoiceSerializer,
    FinanceTransactionSerializer,
)


def _user_is_company_admin(user) -> bool:
    if user.is_staff or user.is_superuser:
        return True
    emp = getattr(user, "employee_profile", None)
    return bool(emp and getattr(emp, "is_company_admin", False))


def _can_write_finance_record(user, owner_id) -> bool:
    if user.is_staff or user.is_superuser:
        return True
    if _user_is_company_admin(user):
        return True
    return owner_id == user.id


class FinanceDashboardView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        tx_qs = FinanceTransaction.objects.all()
        accounts_qs = FinanceAccount.objects.all()
        categories_qs = FinanceCategory.objects.all()

        incomes = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_INCOME).aggregate(total=Sum("amount"))["total"] or Decimal(
            "0"
        )
        expenses = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_EXPENSE).aggregate(total=Sum("amount"))[
            "total"
        ] or Decimal("0")
        transfers = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_TRANSFER).aggregate(total=Sum("amount"))[
            "total"
        ] or Decimal("0")
        opening = accounts_qs.aggregate(total=Sum("opening_balance"))["total"] or Decimal("0")

        return response.Response(
            {
                "accounts_count": accounts_qs.count(),
                "categories_count": categories_qs.count(),
                "transactions_count": tx_qs.count(),
                "income_amount": str(incomes),
                "expense_amount": str(expenses),
                "transfer_amount": str(transfers),
                "opening_balance_total": str(opening),
                "net_cashflow": str((incomes - expenses).quantize(Decimal("0.01"))),
            }
        )


class FinanceAccountViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceAccountSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return FinanceAccount.objects.select_related("owner").all()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        serializer.save()

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        instance.delete()


class FinanceCategoryViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceCategorySerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return FinanceCategory.objects.select_related("owner").all()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        serializer.save()

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        instance.delete()


class FinanceTransactionViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceTransactionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = FinanceTransaction.objects.select_related("owner", "account", "category", "transfer_account").all()
        return self._apply_filters(qs)

    def _apply_filters(self, qs):
        tx_type = self.request.query_params.get("transaction_type")
        account = self.request.query_params.get("account")
        category = self.request.query_params.get("category")
        date_from = self.request.query_params.get("date_from")
        date_to = self.request.query_params.get("date_to")

        if tx_type:
            qs = qs.filter(transaction_type=tx_type)
        if account:
            qs = qs.filter(account_id=account)
        if category:
            qs = qs.filter(category_id=category)
        if date_from:
            qs = qs.filter(booked_on__gte=date_from)
        if date_to:
            qs = qs.filter(booked_on__lte=date_to)
        return qs

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        serializer.save()

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        instance.delete()

    @action(detail=False, methods=["get"], url_path="export-csv")
    def export_csv(self, request):
        qs = self._apply_filters(
            FinanceTransaction.objects.select_related("account", "category", "transfer_account").all()
        ).order_by("-booked_on", "-created_at")
        resp = HttpResponse(content_type="text/csv")
        resp["Content-Disposition"] = 'attachment; filename="finance-transactions.csv"'
        writer = csv.writer(resp)
        writer.writerow(
            [
                "Date",
                "Libelle",
                "Type",
                "Montant",
                "Compte source",
                "Categorie",
                "Compte destinataire",
                "Reference",
                "Notes",
            ]
        )
        for tx in qs:
            writer.writerow(
                [
                    tx.booked_on.isoformat(),
                    tx.title,
                    tx.transaction_type,
                    str(tx.amount),
                    tx.account.name if tx.account else "",
                    tx.category.name if tx.category else "",
                    tx.transfer_account.name if tx.transfer_account else "",
                    tx.reference or "",
                    tx.notes or "",
                ]
            )
        return resp


def _refresh_invoice_totals(invoice: FinanceInvoice) -> None:
    subtotal = Decimal("0")
    tax_total = Decimal("0")
    for line in invoice.lines.all():
        quantity = line.quantity or Decimal("0")
        unit_price = line.unit_price or Decimal("0")
        tax_rate = line.tax_rate or Decimal("0")
        line_subtotal = quantity * unit_price
        line_tax = line_subtotal * tax_rate / Decimal("100")
        subtotal += line_subtotal
        tax_total += line_tax
    total = subtotal + tax_total
    if invoice.paid_amount >= total and total > Decimal("0"):
        status = FinanceInvoice.STATUS_PAID
    elif invoice.paid_amount > Decimal("0"):
        status = FinanceInvoice.STATUS_PARTIAL
    elif invoice.status == FinanceInvoice.STATUS_PAID and invoice.paid_amount == Decimal("0"):
        status = FinanceInvoice.STATUS_ISSUED
    else:
        status = invoice.status
    invoice.subtotal_amount = subtotal.quantize(Decimal("0.01"))
    invoice.tax_amount = tax_total.quantize(Decimal("0.01"))
    invoice.total_amount = total.quantize(Decimal("0.01"))
    invoice.status = status
    invoice.save(update_fields=["subtotal_amount", "tax_amount", "total_amount", "status", "updated_at"])


class FinanceInvoiceViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceInvoiceSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return FinanceInvoice.objects.select_related("owner").prefetch_related("lines").all()

    def perform_create(self, serializer):
        invoice = serializer.save(owner=self.request.user)
        _refresh_invoice_totals(invoice)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        invoice = serializer.save()
        _refresh_invoice_totals(invoice)

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        instance.delete()


class FinanceInvoiceLineViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceInvoiceLineSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = FinanceInvoiceLine.objects.select_related("owner", "invoice", "category").all()
        invoice = self.request.query_params.get("invoice")
        if invoice:
            qs = qs.filter(invoice_id=invoice)
        return qs

    def perform_create(self, serializer):
        invoice = serializer.validated_data.get("invoice")
        if invoice and not _can_write_finance_record(self.request.user, invoice.owner_id):
            raise PermissionDenied()
        line = serializer.save(owner=self.request.user)
        _refresh_invoice_totals(line.invoice)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        line = serializer.save()
        _refresh_invoice_totals(line.invoice)

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        invoice = instance.invoice
        instance.delete()
        _refresh_invoice_totals(invoice)


class FinanceDocumentViewSet(viewsets.ModelViewSet):
    serializer_class = FinanceDocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = FinanceDocument.objects.select_related("owner", "account", "category", "transaction", "invoice").all()
        document_type = self.request.query_params.get("document_type")
        report_scope = self.request.query_params.get("report_scope")
        date_from = self.request.query_params.get("date_from")
        date_to = self.request.query_params.get("date_to")
        if document_type:
            qs = qs.filter(document_type=document_type)
        if report_scope:
            qs = qs.filter(report_scope=report_scope)
        if date_from:
            qs = qs.filter(document_date__gte=date_from)
        if date_to:
            qs = qs.filter(document_date__lte=date_to)
        return qs

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    def perform_update(self, serializer):
        if not _can_write_finance_record(self.request.user, serializer.instance.owner_id):
            raise PermissionDenied()
        serializer.save()

    def perform_destroy(self, instance):
        if not _can_write_finance_record(self.request.user, instance.owner_id):
            raise PermissionDenied()
        instance.delete()


class FinanceReportView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    ALL_REPORT_SECTIONS = {
        "summary",
        "category_breakdown",
        "monthly_trend",
        "annex_transactions",
        "annex_invoices",
        "annex_documents",
    }

    def _get_filters(self, request):
        source = request.query_params if request.method == "GET" else request.data
        return {
            "date_from": source.get("date_from"),
            "date_to": source.get("date_to"),
            "category": source.get("category"),
            "document_type": source.get("document_type"),
            "report_scope": source.get("report_scope"),
            "invoice_status": source.get("invoice_status"),
        }

    def _get_sections(self, request):
        if request.method == "GET":
            raw_sections = request.query_params.getlist("sections")
        else:
            raw_sections = request.data.get("sections", [])

        if isinstance(raw_sections, str):
            raw_sections = [part.strip() for part in raw_sections.split(",") if part.strip()]
        elif raw_sections is None:
            raw_sections = []
        elif not isinstance(raw_sections, list):
            raw_sections = [str(raw_sections)]

        selected = {str(section).strip() for section in raw_sections if str(section).strip() in self.ALL_REPORT_SECTIONS}
        return selected or set(self.ALL_REPORT_SECTIONS)

    def _apply_filters(self, tx_qs, documents_qs, invoices_qs, filters):
        date_from = filters.get("date_from")
        date_to = filters.get("date_to")
        category = filters.get("category")
        document_type = filters.get("document_type")
        report_scope = filters.get("report_scope")
        invoice_status = filters.get("invoice_status")

        if date_from:
            tx_qs = tx_qs.filter(booked_on__gte=date_from)
            documents_qs = documents_qs.filter(document_date__gte=date_from)
            invoices_qs = invoices_qs.filter(issued_on__gte=date_from)
        if date_to:
            tx_qs = tx_qs.filter(booked_on__lte=date_to)
            documents_qs = documents_qs.filter(document_date__lte=date_to)
            invoices_qs = invoices_qs.filter(issued_on__lte=date_to)
        if category:
            tx_qs = tx_qs.filter(category_id=category)
            documents_qs = documents_qs.filter(category_id=category)
        if document_type:
            documents_qs = documents_qs.filter(document_type=document_type)
        if report_scope:
            documents_qs = documents_qs.filter(report_scope=report_scope)
        if invoice_status:
            invoices_qs = invoices_qs.filter(status=invoice_status)
        return tx_qs, documents_qs, invoices_qs

    def _build_report_payload(self, tx_qs, documents_qs, invoices_qs, filters):
        date_from = filters.get("date_from")
        date_to = filters.get("date_to")
        category = filters.get("category")
        document_type = filters.get("document_type")
        report_scope = filters.get("report_scope")
        invoice_status = filters.get("invoice_status")

        income = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_INCOME).aggregate(total=Sum("amount"))["total"] or Decimal("0")
        expense = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_EXPENSE).aggregate(total=Sum("amount"))["total"] or Decimal("0")
        transfer = tx_qs.filter(transaction_type=FinanceTransaction.TYPE_TRANSFER).aggregate(total=Sum("amount"))["total"] or Decimal("0")

        by_category_qs = (
            tx_qs.values("category_id", "category__name")
            .annotate(total=Sum("amount"))
            .order_by("-total")
        )
        by_category = [
            {
                "category_id": row["category_id"],
                "category_name": row["category__name"] or "Sans categorie",
                "total": str(row["total"] or Decimal("0")),
            }
            for row in by_category_qs
        ]

        monthly_qs = (
            tx_qs.annotate(month=TruncMonth("booked_on"))
            .values("month", "transaction_type")
            .annotate(total=Sum("amount"))
            .order_by("month")
        )
        monthly = []
        for row in monthly_qs:
            month_value = row["month"]
            # TruncMonth can return either date or datetime depending on backend.
            if month_value is None:
                month_iso = None
            elif hasattr(month_value, "date"):
                month_iso = month_value.date().isoformat()
            else:
                month_iso = month_value.isoformat()
            monthly.append(
                {
                    "month": month_iso,
                    "transaction_type": row["transaction_type"],
                    "total": str(row["total"] or Decimal("0")),
                }
            )

        invoices_summary = invoices_qs.aggregate(
            billed_total=Sum("total_amount"),
            paid_total=Sum("paid_amount"),
        )
        billed_total = invoices_summary.get("billed_total") or Decimal("0")
        paid_total = invoices_summary.get("paid_total") or Decimal("0")
        outstanding_total = (billed_total - paid_total).quantize(Decimal("0.01"))

        documents_summary = documents_qs.aggregate(total_amount=Sum("amount"))
        documents_total = documents_summary.get("total_amount") or Decimal("0")

        return {
            "filters": {
                "date_from": date_from,
                "date_to": date_to,
                "category": category,
                "document_type": document_type,
                "report_scope": report_scope,
                "invoice_status": invoice_status,
            },
            "summary": {
                "income_amount": str(income),
                "expense_amount": str(expense),
                "transfer_amount": str(transfer),
                "net_cashflow": str((income - expense).quantize(Decimal("0.01"))),
            },
            "invoice_summary": {
                "count": invoices_qs.count(),
                "billed_total": str(billed_total),
                "paid_total": str(paid_total),
                "outstanding_total": str(outstanding_total),
            },
            "document_summary": {
                "count": documents_qs.count(),
                "total_amount": str(documents_total),
            },
            "by_category": by_category,
            "monthly": monthly,
        }

    def _get_or_create_folder(self, name: str, parent, user):
        folder = DocumentFolder.objects.filter(name=name, parent=parent).first()
        if folder:
            return folder
        return DocumentFolder.objects.create(name=name, parent=parent, created_by=user)

    def _build_docx(self, tenant_name: str, payload: dict, transactions, documents, invoices, sections):
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        doc = DocxDocument()
        f = payload["filters"]

        # Style global pour un rendu plus premium
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)
        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Rapport financier detaille")
        run.bold = True
        run.font.size = Pt(20)

        company_line = doc.add_paragraph()
        company_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_line.add_run(f"Entreprise: {tenant_name}")
        company_run.bold = True
        company_run.font.size = Pt(14)

        # Important: ne pas laisser "aujourd hui" (rapport lu plus tard => confusion)
        date_from_display = f.get("date_from") or "debut"
        date_to_display = f.get("date_to") or datetime.now().date().isoformat()

        doc.add_paragraph(
            f"Periode: {date_from_display} -> {date_to_display} | "
            f"Categorie: {f.get('category') or 'toutes'} | Type piece: {f.get('document_type') or 'tous'} | "
            f"Scope: {f.get('report_scope') or 'tous'} | Statut facture: {f.get('invoice_status') or 'tous'}"
        )
        doc.add_paragraph(f"Genere le: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        def _safe_float(value):
            try:
                return float(value)
            except Exception:
                return 0.0

        def _render_category_pie_png():
            # Pie chart: répartition par catégorie
            import matplotlib.pyplot as plt

            by_cat = payload.get("by_category") or []
            values = [_safe_float(row.get("total")) for row in by_cat]

            # Couleurs depuis les catégories (si dispo)
            ids = [row.get("category_id") for row in by_cat if row.get("category_id")]
            cats = FinanceCategory.objects.filter(id__in=ids).values("id", "color")
            color_map = {str(c["id"]): (c.get("color") or "").strip() for c in cats}

            default_colors = ["#185FA5", "#0F6E56", "#B91C1C", "#7C3AED", "#D97706", "#0E7490", "#9333EA", "#475569"]
            colors = []
            for idx, row in enumerate(by_cat):
                cid = row.get("category_id")
                raw = color_map.get(str(cid)) if cid else ""
                c = raw or default_colors[idx % len(default_colors)]
                if c and not c.startswith("#"):
                    c = f"#{c}"
                colors.append(c)

            fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=160)
            ax.pie(
                values,
                colors=colors,
                startangle=90,
                counterclock=False,
                wedgeprops={"linewidth": 1, "edgecolor": "white"},
            )
            ax.set_title("Répartition par catégorie", fontsize=12, fontweight="bold")
            ax.axis("equal")

            buf = BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight")
            plt.close(fig)
            buf.seek(0)
            return buf

        def _render_monthly_bar_png():
            # Bar chart: double bar revenu/dépense
            import matplotlib.pyplot as plt
            import numpy as np

            monthly = payload.get("monthly") or []
            grouped = {}
            for row in monthly:
                month = row.get("month") or ""
                tx_type = row.get("transaction_type") or "income"
                grouped.setdefault(month, {"income": 0.0, "expense": 0.0})
                if tx_type == "income":
                    grouped[month]["income"] += _safe_float(row.get("total"))
                elif tx_type == "expense":
                    grouped[month]["expense"] += _safe_float(row.get("total"))

            months = sorted(grouped.keys())
            incomes = [grouped[m]["income"] for m in months]
            expenses = [grouped[m]["expense"] for m in months]

            fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=160)
            x = np.arange(len(months)) if months else np.array([0])
            width = 0.36
            ax.bar(x - width / 2, incomes or [0], width, color="#0F6E56", label="Revenus")
            ax.bar(x + width / 2, expenses or [0], width, color="#B91C1C", label="Dépenses")

            ax.set_xticks(x)
            ax.set_xticklabels([m[:7] if m else "" for m in months] or ["-"], rotation=45, ha="right", fontsize=8)
            ax.set_ylabel("Montant")
            ax.set_title("Tendance mensuelle", fontsize=12, fontweight="bold")
            ax.legend(fontsize=9, frameon=False)
            fig.tight_layout()

            buf = BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight")
            plt.close(fig)
            buf.seek(0)
            return buf

        if "summary" in sections:
            doc.add_heading("Synthese", level=1)
            summary = payload["summary"]
            inv_sum = payload["invoice_summary"]
            doc_sum = payload["document_summary"]
            doc.add_paragraph(f"Revenus: {summary['income_amount']}")
            doc.add_paragraph(f"Depenses: {summary['expense_amount']}")
            doc.add_paragraph(f"Transferts: {summary['transfer_amount']}")
            doc.add_paragraph(f"Cashflow net: {summary['net_cashflow']}")
            doc.add_paragraph(
                f"Factures: {inv_sum['count']} | Facture total: {inv_sum['billed_total']} | "
                f"Paye: {inv_sum['paid_total']} | Restant: {inv_sum['outstanding_total']}"
            )
            doc.add_paragraph(f"Pieces comptables: {doc_sum['count']} | Montant total pieces: {doc_sum['total_amount']}")

        if "category_breakdown" in sections:
            doc.add_heading("Repartition par categorie", level=1)
            cat_table = doc.add_table(rows=1, cols=3)
            cat_table.style = "Table Grid"
            cat_table.rows[0].cells[0].text = "Categorie"
            cat_table.rows[0].cells[1].text = "ID categorie"
            cat_table.rows[0].cells[2].text = "Montant"
            for row in payload["by_category"]:
                c = cat_table.add_row().cells
                c[0].text = row["category_name"]
                c[1].text = str(row["category_id"] or "")
                c[2].text = row["total"]

            # Graphe pie
            try:
                doc.add_paragraph()
                doc.add_picture(_render_category_pie_png(), width=Cm(16))
            except Exception:
                pass

        if "monthly_trend" in sections:
            doc.add_heading("Tendance mensuelle", level=1)
            m_table = doc.add_table(rows=1, cols=3)
            m_table.style = "Table Grid"
            m_table.rows[0].cells[0].text = "Mois"
            m_table.rows[0].cells[1].text = "Type"
            m_table.rows[0].cells[2].text = "Montant"
            for row in payload["monthly"]:
                c = m_table.add_row().cells
                c[0].text = str(row["month"] or "")
                c[1].text = row["transaction_type"]
                c[2].text = row["total"]

            # Graphe bar
            try:
                doc.add_paragraph()
                doc.add_picture(_render_monthly_bar_png(), width=Cm(16))
            except Exception:
                pass

        if "annex_transactions" in sections:
            doc.add_page_break()
            doc.add_heading("Annexe A - Transactions", level=1)
            tx_table = doc.add_table(rows=1, cols=8)
            tx_table.style = "Table Grid"
            headers = ["Date", "Libelle", "Type", "Compte", "Categorie", "Montant", "Reference", "Owner"]
            for idx, h in enumerate(headers):
                tx_table.rows[0].cells[idx].text = h
            for tx in transactions:
                c = tx_table.add_row().cells
                c[0].text = tx.booked_on.isoformat()
                c[1].text = tx.title
                c[2].text = tx.transaction_type
                c[3].text = tx.account.name if tx.account else ""
                c[4].text = tx.category.name if tx.category else ""
                c[5].text = str(tx.amount)
                c[6].text = tx.reference or ""
                c[7].text = tx.owner.username if tx.owner else ""

        if "annex_invoices" in sections:
            doc.add_page_break()
            doc.add_heading("Annexe B - Factures et lignes", level=1)
            for inv in invoices:
                doc.add_paragraph(
                    f"{inv.number} | Client: {inv.customer_name} | Emission: {inv.issued_on} | "
                    f"Statut: {inv.status} | Total: {inv.total_amount} | Paye: {inv.paid_amount}"
                )
                il_table = doc.add_table(rows=1, cols=6)
                il_table.style = "Table Grid"
                il_headers = ["Description", "Qt", "PU", "Taxe %", "Total ligne", "Categorie"]
                for idx, h in enumerate(il_headers):
                    il_table.rows[0].cells[idx].text = h
                for line in inv.lines.all():
                    c = il_table.add_row().cells
                    c[0].text = line.description
                    c[1].text = str(line.quantity)
                    c[2].text = str(line.unit_price)
                    c[3].text = str(line.tax_rate)
                    c[4].text = str(line.line_total)
                    c[5].text = line.category.name if line.category else ""

        if "annex_documents" in sections:
            doc.add_page_break()
            doc.add_heading("Annexe C - Pieces comptables", level=1)
            d_table = doc.add_table(rows=1, cols=8)
            d_table.style = "Table Grid"
            d_headers = ["Titre", "Type", "Scope", "Date", "Montant", "Compte", "Categorie", "Facture"]
            for idx, h in enumerate(d_headers):
                d_table.rows[0].cells[idx].text = h
            for item in documents:
                c = d_table.add_row().cells
                c[0].text = item.title
                c[1].text = item.document_type
                c[2].text = item.report_scope
                c[3].text = item.document_date.isoformat()
                c[4].text = str(item.amount)
                c[5].text = item.account.name if item.account else ""
                c[6].text = item.category.name if item.category else ""
                c[7].text = item.invoice.number if item.invoice else ""

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def get(self, request):
        tx_qs = FinanceTransaction.objects.select_related("category")
        documents_qs = FinanceDocument.objects.select_related("category")
        invoices_qs = FinanceInvoice.objects.all()
        filters = self._get_filters(request)
        tx_qs, documents_qs, invoices_qs = self._apply_filters(tx_qs, documents_qs, invoices_qs, filters)
        payload = self._build_report_payload(tx_qs, documents_qs, invoices_qs, filters)
        return response.Response(payload)

    def post(self, request):
        tx_qs = FinanceTransaction.objects.select_related("owner", "account", "category")
        documents_qs = FinanceDocument.objects.select_related("account", "category", "invoice")
        invoices_qs = FinanceInvoice.objects.select_related("owner").prefetch_related("lines", "lines__category")
        filters = self._get_filters(request)
        sections = self._get_sections(request)
        tx_qs, documents_qs, invoices_qs = self._apply_filters(tx_qs, documents_qs, invoices_qs, filters)
        payload = self._build_report_payload(tx_qs, documents_qs, invoices_qs, filters)

        tenant = getattr(request, "tenant", None)
        tenant_name = getattr(tenant, "name", "Entreprise")
        content = self._build_docx(
            tenant_name=tenant_name,
            payload=payload,
            transactions=tx_qs.order_by("-booked_on", "-created_at"),
            documents=documents_qs.order_by("-document_date", "-created_at"),
            invoices=invoices_qs.order_by("-issued_on", "-created_at"),
            sections=sections,
        )

        finance_folder = self._get_or_create_folder("Finance", None, request.user)
        reports_folder = self._get_or_create_folder("Rapport financier", finance_folder, request.user)

        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"rapport-financier-{stamp}.docx"
        title = f"Rapport financier {stamp}"

        ged_doc = Document(
            title=title,
            description="Rapport financier detaille genere automatiquement depuis le module Finance.",
            folder=reports_folder,
            original_filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size=len(content),
            uploaded_by=request.user,
        )
        ged_doc.file.save(filename, ContentFile(content), save=False)
        ged_doc.save()

        return response.Response(
            {
                "message": "Rapport financier genere et enregistre dans GED/Finance/Rapport financier.",
                "document_id": str(ged_doc.id),
                "document_title": ged_doc.title,
                "folder_id": str(reports_folder.id),
                "folder_name": reports_folder.name,
                "sections": sorted(sections),
            },
            status=201,
        )

